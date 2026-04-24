"""
Multi-format document loader for Amicus.
Supports: .md, .docx, .pdf, .ipynb, .py (structured legal data)
"""

import importlib.util
import json
import os
import re
from pathlib import Path
from typing import Optional

from langchain_core.documents import Document

TOPIC_MAP = {
    "python": "python",
    "statistics": "statistics",
    "nlp": "nlp",
    "legal": "legal_knowledge",
}

JURISDICTION_LABELS = {
    "CN": "China",
    "UK": "United Kingdom",
}

COMMON_MOJIBAKE_REPLACEMENTS = {
    "â€™": "'",
    "â€œ": '"',
    "â€\x9d": '"',
    "â€˜": "'",
    "â€“": "-",
    "â€”": "-",
    "â€¦": "...",
    "Â£": "£",
    "Â¥": "¥",
    "Â": "",
}

LEGAL_FIELD_LABELS = {
    "title": "Title",
    "name": "Case",
    "text": "Text",
    "source": "Authority",
    "court": "Court",
    "year": "Year",
    "dispute": "Dispute",
    "court_reasoning": "Court Reasoning",
    "final_split": "Outcome",
    "key_holding": "Key Holding",
    "quote": "Quoted Passage",
    "quote_attribution": "Quoted Judge",
    "financial_outcome": "Financial Outcome",
    "ratio_decidendi": "Ratio",
    "tags": "Tags",
}

LEGAL_FIELD_ORDER = (
    "title",
    "name",
    "source",
    "court",
    "year",
    "text",
    "key_holding",
    "dispute",
    "court_reasoning",
    "final_split",
    "financial_outcome",
    "ratio_decidendi",
    "quote",
    "quote_attribution",
    "tags",
)


def _infer_topic(filepath: str) -> str:
    """Infer topic from the subdirectory name."""
    parts = Path(filepath).parts
    for part in parts:
        if part in TOPIC_MAP:
            return TOPIC_MAP[part]
    return "general"


def _repair_common_mojibake(text: str) -> str:
    repaired = str(text or "").replace("\ufeff", "").replace("\xa0", " ")
    for broken, fixed in COMMON_MOJIBAKE_REPLACEMENTS.items():
        repaired = repaired.replace(broken, fixed)
    return repaired


def _clean_inline_text(text: str) -> str:
    return re.sub(r"\s+", " ", _repair_common_mojibake(text)).strip()


def _humanize_identifier(value: str) -> str:
    tokens = re.split(r"[\s_-]+", Path(value).stem)
    cleaned_tokens = []
    for token in tokens:
        if not token:
            continue
        lowered = token.lower()
        if lowered in {"nlp", "uk", "cn", "etl"}:
            cleaned_tokens.append(token.upper())
        elif lowered == "py4e":
            cleaned_tokens.append("PY4E")
        elif token.isdigit():
            cleaned_tokens.append(token)
        else:
            cleaned_tokens.append(token.capitalize())
    return " ".join(cleaned_tokens) or "Reference"


def _truncate_label(text: str, *, limit: int = 96) -> str:
    if len(text) <= limit:
        return text
    return text[: limit - 3].rstrip() + "..."


def _looks_like_heading_noise(text: str) -> bool:
    lowered = text.lower()
    if not lowered:
        return True
    if lowered.startswith(("xml version=", "<?xml", "function ", "var ")):
        return True
    if len(re.findall(r"[a-z]", lowered)) < 3:
        return True
    return bool(re.search(r"[{}<>]{2,}", text))


def _extract_section_title(text: str, fallback: str) -> str:
    lines = [_clean_inline_text(line) for line in (text or "").splitlines()]
    heading_candidates = []
    general_candidates = []

    for raw_line in lines:
        if not raw_line:
            continue
        stripped = re.sub(r"^#{1,6}\s*", "", raw_line).strip()
        if not stripped:
            continue
        if raw_line.startswith("#"):
            heading_candidates.append(stripped)
        general_candidates.append(stripped)

    for candidate in heading_candidates + general_candidates:
        if _looks_like_heading_noise(candidate):
            continue
        return _truncate_label(candidate)

    return _truncate_label(fallback)


def _clean_legal_value(value) -> str:
    if isinstance(value, list):
        return ", ".join(_clean_inline_text(item) for item in value if _clean_inline_text(item))
    return _clean_inline_text(value)


def _legal_material_label(category: str) -> str:
    if category == "Cases":
        return "Case"
    if category == "Statutes":
        return "Statute"
    return category.rstrip("s")


def _legal_source_label(category: str, article_id: str, article_data: dict) -> str:
    preferred_keys = ("name", "source", "title") if category == "Cases" else ("source", "title", "name")
    for key in preferred_keys:
        value = _clean_legal_value(article_data.get(key))
        if value:
            return value
    return _humanize_identifier(article_id)


def _format_legal_record(jurisdiction: str, category: str, article_id: str, article_data: dict) -> str:
    lines = [
        f"Jurisdiction: {JURISDICTION_LABELS.get(jurisdiction, jurisdiction)}",
        f"Material Type: {_legal_material_label(category)}",
        f"Record ID: {article_id}",
    ]
    seen = set()

    for key in LEGAL_FIELD_ORDER:
        value = article_data.get(key)
        cleaned = _clean_legal_value(value)
        if not cleaned:
            continue
        lines.append(f"{LEGAL_FIELD_LABELS.get(key, key.replace('_', ' ').title())}: {cleaned}")
        seen.add(key)

    for key, value in article_data.items():
        if key in seen:
            continue
        cleaned = _clean_legal_value(value)
        if not cleaned:
            continue
        lines.append(f"{LEGAL_FIELD_LABELS.get(key, key.replace('_', ' ').title())}: {cleaned}")

    return "\n".join(lines)


def load_markdown(filepath: str) -> list[Document]:
    """Load a Markdown file, splitting by ## headers into sections."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()

        topic = _infer_topic(filepath)
        fallback_title = _humanize_identifier(os.path.basename(filepath))
        sections = content.split("\n## ")
        documents = []

        for i, section in enumerate(sections):
            text = section.strip()
            if not text:
                continue
            if i > 0:
                text = "## " + text

            title_line = _extract_section_title(text, fallback_title)
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": os.path.basename(filepath),
                        "format": "md",
                        "section_title": title_line,
                        "document_title": fallback_title,
                        "topic": topic,
                    },
                )
            )
        return documents
    except Exception as e:
        print(f"[ERROR] Failed to load markdown {filepath}: {e}")
        return []


def load_docx(filepath: str) -> list[Document]:
    """Load a Word document, extracting each non-empty paragraph."""
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(filepath)
        topic = _infer_topic(filepath)
        document_title = _humanize_identifier(os.path.basename(filepath))
        documents = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": os.path.basename(filepath),
                        "format": "docx",
                        "paragraph_index": i + 1,
                        "document_title": document_title,
                        "topic": topic,
                    },
                )
            )
        return documents
    except Exception as e:
        print(f"[ERROR] Failed to load docx {filepath}: {e}")
        return []


def clean_pdf_text(text: str) -> str:
    """Clean raw PDF-extracted text to improve readability and chunk quality."""
    lines = text.split('\n')
    cleaned_lines = []

    for line in lines:
        stripped = line.strip()
        # Remove isolated single characters (formula residue like "T", "t", "x")
        if len(stripped) <= 2 and not stripped.endswith('.') and not stripped.isdigit():
            continue
        # Remove lines that are just punctuation or whitespace
        if not stripped or re.match(r'^[.\-_=]+$', stripped):
            continue
        cleaned_lines.append(stripped)

    # Rejoin and fix broken line wraps (line ending without sentence-ending punctuation)
    result = ''
    for i, line in enumerate(cleaned_lines):
        if result and not result.endswith(('.', '!', '?', ':', ';', '"', ')')):
            result += ' ' + line
        else:
            if result:
                result += '\n' + line
            else:
                result = line

    # Normalize multiple spaces
    result = re.sub(r' {2,}', ' ', result)
    return result.strip()


def load_pdf(filepath: str) -> list[Document]:
    """Load a PDF file page by page using pdfplumber."""
    try:
        import pdfplumber

        topic = _infer_topic(filepath)
        document_title = _humanize_identifier(os.path.basename(filepath))
        documents = []

        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    raw_text = page.extract_text()
                    if raw_text and raw_text.strip():
                        text = clean_pdf_text(raw_text)
                        if not text:
                            continue
                        documents.append(
                            Document(
                                page_content=text,
                                metadata={
                                    "source": os.path.basename(filepath),
                                    "format": "pdf",
                                    "page_number": page_num,
                                    "document_title": document_title,
                                    "topic": topic,
                                },
                            )
                        )
                except Exception as e:
                    print(f"[WARN] Skipping page {page_num} of {filepath}: {e}")
                    continue
        return documents
    except Exception as e:
        print(f"[ERROR] Failed to load PDF {filepath}: {e}")
        return []


def load_ipynb(filepath: str) -> list[Document]:
    """Load a Jupyter Notebook, extracting markdown and code cells."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            notebook = json.load(f)

        topic = _infer_topic(filepath)
        document_title = _humanize_identifier(os.path.basename(filepath))
        documents = []
        cells = notebook.get("cells", [])

        for i, cell in enumerate(cells):
            cell_type = cell.get("cell_type", "unknown")
            source_lines = cell.get("source", [])
            text = "".join(source_lines).strip()

            if not text:
                continue

            if cell_type == "code":
                text = f"[Code Example]\n{text}"

            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": os.path.basename(filepath),
                        "format": "ipynb",
                        "cell_type": cell_type,
                        "cell_index": i + 1,
                        "cell_title": _extract_section_title(text, f"Cell {i + 1}"),
                        "document_title": document_title,
                        "topic": topic,
                    },
                )
            )
        return documents
    except Exception as e:
        print(f"[ERROR] Failed to load notebook {filepath}: {e}")
        return []


def load_legal_py(filepath: str) -> list[Document]:
    """Load structured legal data from a Python module."""
    try:
        spec = importlib.util.spec_from_file_location("legal_data", filepath)
        if spec is None or spec.loader is None:
            print(f"[ERROR] Failed to load legal data spec for {filepath}")
            return []
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)

        knowledge_base = getattr(module, "LEGAL_KNOWLEDGE_BASE", None)
        if knowledge_base is None:
            print(f"[WARN] No LEGAL_KNOWLEDGE_BASE found in {filepath}")
            return []

        documents = []
        for jurisdiction, categories in knowledge_base.items():
            for category, articles in categories.items():
                for article_id, article_data in articles.items():
                    if isinstance(article_data, dict):
                        text = _format_legal_record(jurisdiction, category, article_id, article_data)
                        record_title = (
                            _clean_legal_value(article_data.get("title"))
                            or _clean_legal_value(article_data.get("name"))
                            or _clean_legal_value(article_data.get("source"))
                            or _humanize_identifier(article_id)
                        )
                        source_label = _legal_source_label(category, article_id, article_data)
                        authority = _clean_legal_value(article_data.get("source"))
                        court = _clean_legal_value(article_data.get("court"))
                        year = article_data.get("year")
                    else:
                        text = _clean_inline_text(article_data)
                        record_title = _humanize_identifier(article_id)
                        source_label = record_title
                        authority = ""
                        court = ""
                        year = None

                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": source_label,
                                "format": "py",
                                "jurisdiction": jurisdiction,
                                "jurisdiction_name": JURISDICTION_LABELS.get(jurisdiction, jurisdiction),
                                "category": category,
                                "article_id": article_id,
                                "record_title": record_title,
                                "authority": authority,
                                "court": court,
                                "year": year,
                                "topic": "legal_knowledge",
                            },
                        )
                    )
        return documents
    except Exception as e:
        print(f"[ERROR] Failed to load legal data {filepath}: {e}")
        return []


LOADERS = {
    ".md": load_markdown,
    ".docx": load_docx,
    ".pdf": load_pdf,
    ".ipynb": load_ipynb,
    ".py": load_legal_py,
}


def load_all_documents(data_dir: str = "data/raw") -> list[Document]:
    """
    Walk through data_dir, detect file format, apply correct loader.
    Infer 'topic' from subdirectory name (python/, statistics/, nlp/, legal/).
    Return list of LangChain Document objects.
    """
    all_docs = []
    stats = {"by_format": {}, "by_topic": {}}

    for root, _dirs, files in os.walk(data_dir):
        for filename in sorted(files):
            filepath = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()

            if ext not in LOADERS:
                continue

            # Skip __init__.py or non-data Python files
            if ext == ".py" and "legal_data" not in filename:
                continue

            print(f"Loading: {filepath}")
            loader = LOADERS[ext]
            docs = loader(filepath)
            all_docs.extend(docs)

            fmt = ext.lstrip(".")
            stats["by_format"][fmt] = stats["by_format"].get(fmt, 0) + len(docs)
            for doc in docs:
                t = doc.metadata.get("topic", "unknown")
                stats["by_topic"][t] = stats["by_topic"].get(t, 0) + 1

    print("\n" + "=" * 50)
    print("DOCUMENT LOADING SUMMARY")
    print("=" * 50)
    print(f"Total documents: {len(all_docs)}")
    print("\nBy format:")
    for fmt, count in sorted(stats["by_format"].items()):
        print(f"  .{fmt}: {count}")
    print("\nBy topic:")
    for topic, count in sorted(stats["by_topic"].items()):
        print(f"  {topic}: {count}")

    return all_docs


if __name__ == "__main__":
    docs = load_all_documents()
    if docs:
        print(f"\nSample document (first):")
        print(f"  Content: {docs[0].page_content[:200]}...")
        print(f"  Metadata: {docs[0].metadata}")
